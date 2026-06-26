import pytest

from app.core.context import AgentContext, set_agent_context


pytestmark = pytest.mark.no_infrastructure


@pytest.fixture
def word_context(tmp_path, monkeypatch):
    from docx import Document
    from app.services.ai.tools import document_paths, generated_file_service

    uploads = tmp_path / "uploads"
    uploads.mkdir()
    source = uploads / "letter.docx"
    document = Document()
    document.add_paragraph("Hello {{name}}")
    document.add_paragraph("Second paragraph")
    document.save(source)
    monkeypatch.setattr(document_paths, "get_data_base_dir", lambda: str(tmp_path))
    async def workspace_root():
        return str(tmp_path / "agent_workspaces")
    monkeypatch.setattr(document_paths, "resolve_workspace_root", workspace_root)
    monkeypatch.setattr(generated_file_service, "generated_files_root", lambda: tmp_path / "generated")
    set_agent_context(AgentContext(
        agent_id="agent", agent_name="Agent", user_id=1, conversation_id="conv",
        authorized_attachment_paths=[str(source)],
    ))
    return source


@pytest.mark.asyncio
async def test_word_read_content_is_paginated(word_context):
    from app.services.ai.tools.word_document_tool import word_document_read

    result = await word_document_read.ainvoke({
        "action": "read_content", "path": str(word_context), "start": 1, "limit": 1,
    })

    assert result["data"]["paragraphs"] == ["Second paragraph"]


@pytest.mark.asyncio
async def test_word_replace_text_writes_a_copy(word_context):
    from app.services.ai.tools.word_document_tool import word_document_write

    result = await word_document_write.ainvoke({
        "action": "replace_text", "path": str(word_context),
        "replacements": [{"find": "{{name}}", "replace": "Alice"}],
        "output_filename": "letter_alice.docx",
    })

    assert result["changes"]["replacements"] == 1
    assert result["artifact"]["download_url"]
